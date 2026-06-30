import re
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
from .gemini_helper import generate_text
from .content_generator import (
    GROUND_RULES, section_guidance, trim_to_sentence, _high_word_bound, _crossref_citations,
    strip_section_heading,
)
import io

class DefaultPaperGenerator:
    """Generate default research papers in Word format with balanced structure."""
    
    def __init__(self):
        self.section_word_limits = {
            "Abstract": "150-200",
            "Introduction": "200-250", 
            "Literature Review": "200-250",
            "Methodology": "150-200",
            "Results": "200-250",
            "Discussion": "200-250",
            "Conclusion": "150-200"
        }
    
    def generate_default_paper(self, topic: str, paper_type: str = "empirical", suggestions: dict = None, images: list = None) -> Dict[str, Any]:
        """Generate a balanced, type-specific research paper with advanced realism, academic polish, user suggestions, and images."""
        # Section structures for each type
        type_structures = {
            "empirical": ["Abstract", "Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion"],
            "theoretical": ["Abstract", "Introduction", "Theoretical Framework", "Literature Review", "Analysis", "Discussion", "Conclusion"],
            "review": ["Abstract", "Introduction", "Methodology", "Literature Review", "Analysis", "Discussion", "Conclusion"],
            "comparative": ["Abstract", "Introduction", "Literature Review", "Methodology", "Comparative Analysis", "Discussion", "Conclusion"],
            "case_study": ["Abstract", "Introduction", "Case Background", "Methodology", "Case Analysis", "Discussion", "Conclusion"],
            "analytical": ["Abstract", "Introduction", "Literature Review", "Analytical Framework", "Analysis", "Discussion", "Conclusion"],
            "methodological": ["Abstract", "Introduction", "Literature Review", "Methodology Development", "Validation", "Discussion", "Conclusion"],
            "position": ["Abstract", "Introduction", "Background", "Position Statement", "Supporting Arguments", "Discussion", "Conclusion"],
            "technical": ["Abstract", "Introduction", "Technical Background", "Methods", "Results", "Technical Analysis", "Conclusion"],
            "interdisciplinary": ["Abstract", "Introduction", "Theoretical Integration", "Methodology", "Cross-Disciplinary Analysis", "Discussion", "Conclusion"]
        }
        # Section word limits (default fallback)
        default_word_limits = {
            "Abstract": "150-200", "Introduction": "200-250", "Literature Review": "200-250", "Methodology": "150-200", "Results": "200-250", "Discussion": "200-250", "Conclusion": "150-200",
            "Theoretical Framework": "200-250", "Analysis": "200-250", "Case Background": "200-250", "Case Analysis": "200-250", "Comparative Analysis": "200-250", "Analytical Framework": "200-250", "Methodology Development": "200-250", "Validation": "200-250", "Background": "200-250", "Position Statement": "200-250", "Supporting Arguments": "200-250", "Technical Background": "200-250", "Methods": "200-250", "Technical Analysis": "200-250", "Theoretical Integration": "200-250", "Cross-Disciplinary Analysis": "200-250"
        }
        # A single grounded thesis — no fabricated sample size / statistic /
        # "advanced method" (e.g. network analysis) / "extra angle" (e.g. ESG).
        thesis = generate_text(
            f'In one clear, factual sentence, state the central question or thesis that a '
            f'{paper_type} research paper on "{topic}" would address. Be specific. Do not invent statistics.'
        ).strip()

        # Use type-specific structure
        structure = type_structures.get(paper_type, type_structures["empirical"])
        # Generate paper content
        paper = {
            "topic": topic,
            "paper_type": paper_type,
            "title": self._generate_title(topic),
            "structure": structure,
            "sections": {},
            "citations": _crossref_citations(topic),  # REAL references via CrossRef
            "word_count": 0,
            "thesis": thesis,
            "images": images or []
        }
        for section_name in structure:
            word_limit = default_word_limits.get(section_name, "200-250")
            suggestion_text = ""
            if suggestions:
                # If suggestions is a dict, check for section-specific suggestions
                if isinstance(suggestions, dict) and section_name in suggestions:
                    suggestion_text = suggestions[section_name]
                # If suggestions is a string, use for all sections
                elif isinstance(suggestions, str):
                    suggestion_text = suggestions
            content = self._generate_section(
                section_name, topic, word_limit, thesis, paper_type, suggestion_text
            )
            paper["sections"][section_name] = {
                "content": content,
                "word_count": len(content.split()),
                "word_limit": word_limit
            }
            paper["word_count"] += len(content.split())
        return paper
    
    def _generate_title(self, topic: str) -> str:
        """Generate a clear, professional title."""
        title_prompt = f"""
        Create a clear, professional research paper title for: {topic}
        
        Requirements:
        - Short and descriptive (under 15 words)
        - Include key keywords
        - Professional academic style
        - Avoid jargon when possible
        
        Format: Main topic - Subtitle or focus area
        """
        
        try:
            return generate_text(title_prompt).strip()
        except Exception as e:
            return f"Research on {topic}"
    
    def _generate_section(self, section_name: str, topic: str, word_limit: str,
                          thesis: str = "", paper_type: str = "empirical", suggestion_text: str = "") -> str:
        """Generate a section with fabrication-free, type-appropriate prompting."""
        thesis_line = f'Central question / thesis: {thesis}\n' if thesis else ""
        sugg_line = f'Additional focus requested by the user: {suggestion_text}\n' if suggestion_text else ""
        prompt = (
            f'Write the "{section_name}" section of a {paper_type} research paper on the topic: "{topic}".\n'
            f'{thesis_line}{sugg_line}'
            f'\nWhat this section should contain:\n{section_guidance(section_name)}\n'
            f'\nTarget length: about {word_limit} words.\n\n'
            f'{GROUND_RULES}'
        )
        try:
            content = generate_text(prompt).strip()
        except Exception:
            return f"The {section_name} section could not be generated. Please try again."
        content = strip_section_heading(content, section_name)
        return trim_to_sentence(content, _high_word_bound(word_limit))

    def create_word_document(self, paper: Dict[str, Any]) -> bytes:
        """Create a Word document (.docx) from the paper, inserting images into the correct sections."""
        try:
            doc = docx.Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Helper: get images for a section
            def get_section_images(section):
                return [img for img in paper.get('images', []) if img.get('section') == section]
            
            # Title
            title = doc.add_heading(paper['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Render sections in the paper's OWN (type-specific) order, so every
            # paper type keeps its real structure (Case Analysis, Theoretical
            # Framework, Position Statement, etc.) instead of a fixed empirical list.
            section_order = paper.get('structure') or list(paper['sections'].keys())

            # Add all sections (with images)
            for section in section_order:
                if section in paper['sections']:
                    doc.add_heading(section, level=1)
                    doc.add_paragraph(paper['sections'][section]['content'])
                    # Insert images for this section
                    for img in get_section_images(section):
                        try:
                            doc.add_picture(io.BytesIO(img['image_bytes']), width=Inches(4.5))
                            if img.get('caption'):
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                doc.add_paragraph(img['caption']).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            doc.add_paragraph("[Image could not be loaded]")
                    doc.add_paragraph()
            
            # References
            doc.add_heading('References', level=1)
            for i, citation in enumerate(paper['citations'], 1):
                doc.add_paragraph(f"{i}. {citation['citation']}")
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
            
        except Exception as e:
            # Return a simple text document if Word creation fails. Generic over
            # whatever sections the paper actually has (any paper type).
            lines = [paper.get('title', 'Research Paper'), ""]
            order = paper.get('structure') or list(paper.get('sections', {}).keys())
            for section in order:
                if section in paper.get('sections', {}):
                    lines.append(section)
                    lines.append(paper['sections'][section].get('content', ''))
                    lines.append("")
            lines.append("References")
            for i, citation in enumerate(paper.get('citations', []), 1):
                lines.append(f"{i}. {citation['citation']}")
            return "\n".join(lines).encode('utf-8')

# Lazy global instance — built on first use, reused across reruns.
_default_paper_generator = None


def get_default_paper_generator() -> "DefaultPaperGenerator":
    """Return the shared DefaultPaperGenerator, constructing it on first use."""
    global _default_paper_generator
    if _default_paper_generator is None:
        _default_paper_generator = DefaultPaperGenerator()
    return _default_paper_generator


def __getattr__(name):
    """Backward-compat: ``default_paper_generator.default_paper_generator`` resolves (lazily)."""
    if name == "default_paper_generator":
        return get_default_paper_generator()
    raise AttributeError(f"module {__name__!r} has no attribute {name!r}")


def generate_default_paper(topic: str, paper_type: str = "empirical", suggestions: dict = None, images: list = None) -> Dict[str, Any]:
    """Generate a default research paper."""
    return get_default_paper_generator().generate_default_paper(topic, paper_type, suggestions, images)

def create_word_document(paper: Dict[str, Any]) -> bytes:
    """Create Word document from paper."""
    return get_default_paper_generator().create_word_document(paper)


def text_to_docx(title: str, text: str) -> bytes:
    """Render a heading + body text into a .docx and return its bytes.

    Used to offer a one-click Word download for paraphraser / grammar / section
    outputs (which were previously TXT-only)."""
    # Strip XML-incompatible control chars (keep \t and \n) — python-docx raises
    # ValueError on NUL/\x0b/etc., which would crash the download button.
    safe = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text or "")
    doc = docx.Document()
    if title:
        doc.add_heading(re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", title), level=1)
    for para in safe.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()